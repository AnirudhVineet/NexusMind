"""AI Content Repurposing API — Track H.

POST   /api/content/generate          — enqueue a generation job
GET    /api/content/{id}              — fetch a single generated content row
GET    /api/content                   — list generated content (filterable)
PATCH  /api/content/{id}             — edit content_json (re-runs fidelity)
DELETE /api/content/{id}             — delete row + unlink file
POST   /api/content/{id}/variants    — generate N variant jobs
GET    /api/content/{id}/export      — download as txt / docx / json
GET    /api/content/{id}/image       — download meme PNG
"""
from __future__ import annotations

import json
import uuid
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from fastapi import APIRouter, Depends, Query, Response, status
from fastapi.responses import FileResponse
from pydantic import BaseModel
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.auth.dependencies import get_current_user
from app.core.exceptions import NotFoundError, ValidationError
from app.db.session import get_session
from app.models.generated_content import (
    CONTENT_TYPES,
    SOURCE_TYPES,
    STYLES,
    GeneratedContent,
)
from app.models.user import User

router = APIRouter(prefix="/api", tags=["content"])


# ---------------------------------------------------------------------------
# Schemas
# ---------------------------------------------------------------------------

class GenerateRequest(BaseModel):
    source_type: str
    source_id: uuid.UUID
    content_type: str
    style: str
    platform: str | None = None
    strict_mode: bool = False
    # Reel scripts only — "short" (3-5 beats, <=90s) or "long" (8-14 beats, 5-15min).
    length: str = "short"


class ContentPatch(BaseModel):
    content_json: dict[str, Any]


class ContentOut(BaseModel):
    id: uuid.UUID
    source_type: str
    source_id: uuid.UUID
    content_type: str
    style: str
    content_json: dict[str, Any] | None
    source_chunks: list
    fidelity_flags: list
    file_path: str | None
    status: str
    created_at: str
    edited_at: str | None


# ---------------------------------------------------------------------------
# Serializer
# ---------------------------------------------------------------------------

def _serialize(row: GeneratedContent) -> ContentOut:
    return ContentOut(
        id=row.id,
        source_type=row.source_type,
        source_id=row.source_id,
        content_type=row.content_type,
        style=row.style,
        content_json=row.content_json,
        source_chunks=row.source_chunks or [],
        fidelity_flags=row.fidelity_flags or [],
        file_path=row.file_path,
        status=row.status,
        created_at=row.created_at.isoformat(),
        edited_at=row.edited_at.isoformat() if row.edited_at else None,
    )


async def _get_owned(
    content_id: uuid.UUID, user: User, session: AsyncSession
) -> GeneratedContent:
    row = (
        await session.execute(
            select(GeneratedContent).where(
                GeneratedContent.id == content_id,
                GeneratedContent.user_id == user.id,
            )
        )
    ).scalar_one_or_none()
    if row is None:
        raise NotFoundError("Generated content not found")
    return row


# ---------------------------------------------------------------------------
# Endpoints
# ---------------------------------------------------------------------------

@router.post(
    "/content/generate",
    status_code=status.HTTP_202_ACCEPTED,
    response_model=ContentOut,
)
async def generate_content(
    body: GenerateRequest,
    current_user: User = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
) -> ContentOut:
    if body.source_type not in SOURCE_TYPES:
        raise ValidationError(
            f"source_type must be one of: {', '.join(SOURCE_TYPES)}"
        )
    if body.content_type not in CONTENT_TYPES:
        raise ValidationError(
            f"content_type must be one of: {', '.join(CONTENT_TYPES)}"
        )
    if body.style not in STYLES:
        raise ValidationError(f"style must be one of: {', '.join(STYLES)}")

    row = GeneratedContent(
        user_id=current_user.id,
        source_type=body.source_type,
        source_id=body.source_id,
        content_type=body.content_type,
        style=body.style,
        content_json=None,
        source_chunks=[],
        fidelity_flags=[],
        status="queued",
    )
    session.add(row)
    await session.commit()
    await session.refresh(row)

    # Build kwargs for worker (platform is optional)
    kwargs: dict[str, Any] = {"strict_mode": body.strict_mode}
    if body.platform:
        kwargs["platform"] = body.platform
    if body.content_type == "reel_script" and body.length in ("short", "long"):
        kwargs["length"] = body.length

    from app.workers.content import generate_content_task

    generate_content_task.apply_async(args=[str(row.id)], kwargs=kwargs)

    return _serialize(row)


@router.get("/content/{content_id}", response_model=ContentOut)
async def get_content(
    content_id: uuid.UUID,
    current_user: User = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
) -> ContentOut:
    row = await _get_owned(content_id, current_user, session)
    return _serialize(row)


class ContentListResponse(BaseModel):
    items: list[ContentOut]
    total: int


@router.get("/content", response_model=ContentListResponse)
async def list_content(
    source_id: uuid.UUID | None = Query(default=None),
    content_type: str | None = Query(default=None),
    limit: int = Query(default=20, ge=1, le=200),
    current_user: User = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
) -> ContentListResponse:
    from sqlalchemy import func

    base_filters = [GeneratedContent.user_id == current_user.id]
    if source_id is not None:
        base_filters.append(GeneratedContent.source_id == source_id)
    if content_type is not None:
        base_filters.append(GeneratedContent.content_type == content_type)

    total = (
        await session.execute(
            select(func.count(GeneratedContent.id)).where(*base_filters)
        )
    ).scalar_one()

    stmt = (
        select(GeneratedContent)
        .where(*base_filters)
        .order_by(GeneratedContent.created_at.desc())
        .limit(limit)
    )

    rows = (await session.execute(stmt)).scalars().all()
    return ContentListResponse(items=[_serialize(r) for r in rows], total=int(total or 0))


@router.patch("/content/{content_id}", response_model=ContentOut)
async def patch_content(
    content_id: uuid.UUID,
    body: ContentPatch,
    current_user: User = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
) -> ContentOut:
    row = await _get_owned(content_id, current_user, session)
    row.content_json = body.content_json
    row.edited_at = datetime.now(timezone.utc)

    # Re-run fidelity checks on string values in the edited content_json
    claims: list[tuple[str, str]] = []
    _collect_strings(body.content_json, claims)
    if claims and row.source_chunks:
        flags = await _run_fidelity_checks_async(claims, list(row.source_chunks))
        row.fidelity_flags = flags

    await session.commit()
    await session.refresh(row)
    return _serialize(row)


async def _run_fidelity_checks_async(
    claims: list[tuple[str, str]], source_chunks: list[dict]
) -> list[dict]:
    from app.services.content_engine import _run_fidelity_checks

    return await _run_fidelity_checks(claims, source_chunks)


def _collect_strings(
    obj: Any, out: list[tuple[str, str]], prefix: str = "field"
) -> None:
    """Recursively collect non-empty string values from a nested dict/list."""
    if isinstance(obj, str) and obj.strip():
        out.append((prefix, obj))
    elif isinstance(obj, dict):
        for k, v in obj.items():
            _collect_strings(v, out, prefix=f"{prefix}.{k}")
    elif isinstance(obj, list):
        for i, item in enumerate(obj):
            _collect_strings(item, out, prefix=f"{prefix}[{i}]")


@router.delete(
    "/content/{content_id}",
    status_code=status.HTTP_204_NO_CONTENT,
    response_model=None,
)
async def delete_content(
    content_id: uuid.UUID,
    current_user: User = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
) -> None:
    row = await _get_owned(content_id, current_user, session)

    # Unlink file if present
    if row.file_path:
        try:
            Path(row.file_path).unlink(missing_ok=True)
        except Exception:
            pass

    await session.delete(row)
    await session.commit()


@router.post(
    "/content/{content_id}/variants",
    status_code=status.HTTP_202_ACCEPTED,
    response_model=list[ContentOut],
)
async def create_variants(
    content_id: uuid.UUID,
    n: int = Query(default=3, ge=1, le=10),
    current_user: User = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
) -> list[ContentOut]:
    original = await _get_owned(content_id, current_user, session)

    from app.workers.content import generate_content_task

    new_rows: list[GeneratedContent] = []
    for _ in range(n):
        variant = GeneratedContent(
            user_id=current_user.id,
            source_type=original.source_type,
            source_id=original.source_id,
            content_type=original.content_type,
            style=original.style,
            content_json=None,
            source_chunks=[],
            fidelity_flags=[],
            status="queued",
        )
        session.add(variant)
        new_rows.append(variant)

    await session.commit()
    for v in new_rows:
        await session.refresh(v)
        generate_content_task.apply_async(args=[str(v.id)])

    return [_serialize(v) for v in new_rows]


@router.get("/content/{content_id}/export")
async def export_content(
    content_id: uuid.UUID,
    format: str = Query(default="json", pattern="^(txt|docx|json)$"),
    current_user: User = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
) -> Response:
    row = await _get_owned(content_id, current_user, session)

    if row.content_json is None:
        raise ValidationError("Content is not yet generated")

    if format == "json":
        data = json.dumps(row.content_json, indent=2, ensure_ascii=False).encode("utf-8")
        return Response(
            content=data,
            media_type="application/json",
            headers={
                "Content-Disposition": f'attachment; filename="content_{content_id}.json"'
            },
        )

    if format == "txt":
        text = _content_to_text(row)
        return Response(
            content=text.encode("utf-8"),
            media_type="text/plain",
            headers={
                "Content-Disposition": f'attachment; filename="content_{content_id}.txt"'
            },
        )

    if format == "docx":
        docx_bytes = _content_to_docx(row)
        return Response(
            content=docx_bytes,
            media_type=(
                "application/vnd.openxmlformats-officedocument"
                ".wordprocessingml.document"
            ),
            headers={
                "Content-Disposition": f'attachment; filename="content_{content_id}.docx"'
            },
        )

    raise ValidationError(f"Unsupported format: {format}")


def _content_to_text(row: GeneratedContent) -> str:
    lines = [
        f"Content Type: {row.content_type}",
        f"Style: {row.style}",
        f"Generated: {row.created_at.isoformat()}",
        "",
    ]
    if row.content_json:
        lines.append(json.dumps(row.content_json, indent=2, ensure_ascii=False))
    return "\n".join(lines)


def _content_to_docx(row: GeneratedContent) -> bytes:
    try:
        import io

        from docx import Document as DocxDocument

        doc = DocxDocument()
        doc.add_heading(f"Generated Content: {row.content_type}", level=1)
        doc.add_paragraph(f"Style: {row.style}")
        doc.add_paragraph(f"Generated: {row.created_at.isoformat()}")
        doc.add_paragraph("")
        if row.content_json:
            doc.add_paragraph(json.dumps(row.content_json, indent=2, ensure_ascii=False))
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
    except ImportError:
        # python-docx not available — return plain text as bytes
        return _content_to_text(row).encode("utf-8")


@router.get("/content/{content_id}/image")
async def get_content_image(
    content_id: uuid.UUID,
    current_user: User = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
) -> FileResponse:
    row = await _get_owned(content_id, current_user, session)

    if row.content_type != "meme":
        raise ValidationError("Image export is only available for meme content type")
    if not row.file_path:
        raise NotFoundError("No rendered image for this meme")

    img_path = Path(row.file_path)
    if not img_path.exists():
        raise NotFoundError("Image file not found on disk")

    return FileResponse(
        path=str(img_path),
        media_type="image/png",
        filename=f"meme_{content_id}.png",
    )
